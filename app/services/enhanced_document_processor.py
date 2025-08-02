"""
Enhanced Document Processing Service
Comprehensive document processing with structure preservation, preprocessing, and reconstruction.
"""
import os
import re
import uuid
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple, BinaryIO
from dataclasses import dataclass
import asyncio

import docx
from docx import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import pdfplumber
import textstat
from sentence_transformers import SentenceTransformer
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from sqlalchemy.orm import Session
import regex as re_advanced

from app.core.config import settings
from app.models.document import Document, DocumentStatus
from app.core.database import get_db

logger = logging.getLogger(__name__)

# Download required NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
except:
    logger.warning("Failed to download NLTK data")


@dataclass
class DocumentStructure:
    """Structure information for a document."""
    headings: List[Dict[str, Any]]
    paragraphs: List[Dict[str, Any]]
    citations: List[str]
    academic_terms: List[str]
    metadata: Dict[str, Any]


@dataclass
class TextSegment:
    """Text segment with metadata."""
    text: str
    segment_type: str  # 'heading', 'paragraph', 'citation', 'list_item'
    level: int  # For headings
    style: Optional[str]
    position: int
    formatting: Dict[str, Any]


class EnhancedDocumentProcessor:
    """Enhanced service for comprehensive document processing."""
    
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        # Academic terms patterns (Indonesian and English)
        self.academic_patterns = [
            r'\b(?:penelitian|riset|analisis|metode|metodologi)\b',
            r'\b(?:research|analysis|method|methodology|study)\b',
            r'\b(?:hipotesis|teori|konsep|framework)\b',
            r'\b(?:hypothesis|theory|concept|framework)\b',
            r'\b(?:variabel|sampel|populasi|data)\b',
            r'\b(?:variable|sample|population|dataset)\b',
        ]
        
        # Citation patterns
        self.citation_patterns = [
            r'\([^)]*\d{4}[^)]*\)',  # (Author, 2023)
            r'\[[^\]]*\d{4}[^\]]*\]',  # [Author, 2023]
            r'\d{1,3}\.\s*[A-Z][^.]+\.\s*\(\d{4}\)',  # Bibliography format
            r'(?:et al\.|dkk\.)',  # et al. or dkk.
        ]
        
        # Initialize sentence transformer for similarity
        try:
            self.sentence_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
        except Exception as e:
            logger.warning(f"Failed to load sentence transformer: {e}")
            self.sentence_model = None
    
    async def process_document_enhanced(
        self, 
        file: BinaryIO, 
        filename: str,
        db: Session,
        chapter: Optional[str] = None,
        preserve_structure: bool = True,
        extract_academic_terms: bool = True
    ) -> Document:
        """
        Enhanced document processing with structure preservation.
        
        Args:
            file: File-like object containing the document
            filename: Original filename
            db: Database session
            chapter: Chapter designation
            preserve_structure: Whether to preserve document structure
            extract_academic_terms: Whether to extract academic terminology
            
        Returns:
            Document: Created document record with enhanced metadata
        """
        # Validate file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in settings.allowed_extensions:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Generate unique file path
        file_id = str(uuid.uuid4())
        safe_filename = f"{file_id}{file_ext}"
        file_path = self.upload_dir / safe_filename
        
        try:
            # Save file to disk
            with open(file_path, "wb") as f:
                content = await file.read() if hasattr(file, 'read') else file.read()
                f.write(content)
            
            # Extract text with structure
            if preserve_structure:
                text_content, document_structure = await self._extract_text_with_structure(
                    file_path, file_ext
                )
            else:
                text_content = await self._extract_text_basic(file_path, file_ext)
                document_structure = None
            
            # Preprocess text
            preprocessed_text = await self._preprocess_text(
                text_content, 
                extract_academic_terms=extract_academic_terms
            )
            
            # Calculate quality metrics
            quality_metrics = self._calculate_quality_metrics(text_content)
            
            # Prepare metadata
            metadata = {
                "original_filename": filename,
                "file_size": os.path.getsize(file_path),
                "quality_metrics": quality_metrics,
                "processing_options": {
                    "preserve_structure": preserve_structure,
                    "extract_academic_terms": extract_academic_terms
                }
            }
            
            if document_structure:
                metadata["document_structure"] = {
                    "headings_count": len(document_structure.headings),
                    "paragraphs_count": len(document_structure.paragraphs),
                    "citations_count": len(document_structure.citations),
                    "academic_terms_count": len(document_structure.academic_terms)
                }
            
            # Create document record
            document = Document(
                filename=filename,
                chapter=chapter,
                original_content=text_content,
                file_path=str(file_path),
                status=DocumentStatus.PENDING,
                document_metadata=metadata
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            logger.info(f"Enhanced document processing completed: {document.id}")
            return document
            
        except Exception as e:
            # Clean up file if processing failed
            if file_path.exists():
                file_path.unlink()
            raise IOError(f"Failed to process document: {str(e)}")
    
    async def _extract_text_with_structure(
        self, 
        file_path: Path, 
        file_ext: str
    ) -> Tuple[str, DocumentStructure]:
        """Extract text while preserving document structure."""
        
        if file_ext == ".pdf":
            return await self._extract_pdf_with_structure(file_path)
        elif file_ext == ".docx":
            return await self._extract_docx_with_structure(file_path)
        elif file_ext == ".txt":
            return await self._extract_txt_with_structure(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")
    
    async def _extract_pdf_with_structure(self, file_path: Path) -> Tuple[str, DocumentStructure]:
        """Extract PDF with structure preservation using pdfplumber."""
        text_segments = []
        citations = []
        academic_terms = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text with layout information
                    text = page.extract_text()
                    if not text:
                        continue
                    
                    # Process text line by line to identify structure
                    lines = text.split('\n')
                    for line_num, line in enumerate(lines):
                        line = line.strip()
                        if not line:
                            continue
                        
                        segment = TextSegment(
                            text=line,
                            segment_type=self._classify_text_segment(line),
                            level=0,
                            style=None,
                            position=page_num * 1000 + line_num,
                            formatting={}
                        )
                        text_segments.append(segment)
                        
                        # Extract citations
                        citations.extend(self._extract_citations(line))
                        
                        # Extract academic terms
                        academic_terms.extend(self._extract_academic_terms(line))
        
        except Exception as e:
            logger.error(f"PDF structure extraction failed: {e}")
            # Fallback to basic extraction
            text_content = await self._extract_pdf_text_basic(file_path)
            return text_content, DocumentStructure([], [], [], [], {})
        
        # Combine text segments
        full_text = '\n\n'.join([seg.text for seg in text_segments])
        
        # Create structure
        headings = [
            {"text": seg.text, "level": seg.level, "position": seg.position}
            for seg in text_segments if seg.segment_type == "heading"
        ]
        paragraphs = [
            {"text": seg.text, "position": seg.position}
            for seg in text_segments if seg.segment_type == "paragraph"
        ]
        
        structure = DocumentStructure(
            headings=headings,
            paragraphs=paragraphs,
            citations=list(set(citations)),
            academic_terms=list(set(academic_terms)),
            metadata={"total_pages": len(text_segments)}
        )
        
        return full_text, structure
    
    async def _extract_docx_with_structure(self, file_path: Path) -> Tuple[str, DocumentStructure]:
        """Extract DOCX with full structure preservation."""
        text_segments = []
        citations = []
        academic_terms = []
        
        try:
            doc = DocxDocument(file_path)
            
            for para_num, paragraph in enumerate(doc.paragraphs):
                if not paragraph.text.strip():
                    continue
                
                # Determine paragraph style and level
                style_name = paragraph.style.name if paragraph.style else "Normal"
                is_heading = "Heading" in style_name
                heading_level = 0
                
                if is_heading:
                    try:
                        heading_level = int(style_name.split()[-1])
                    except:
                        heading_level = 1
                
                # Extract formatting information
                formatting = {
                    "bold": paragraph.runs[0].bold if paragraph.runs else False,
                    "italic": paragraph.runs[0].italic if paragraph.runs else False,
                    "alignment": str(paragraph.alignment) if paragraph.alignment else "left"
                }
                
                segment = TextSegment(
                    text=paragraph.text,
                    segment_type="heading" if is_heading else "paragraph",
                    level=heading_level,
                    style=style_name,
                    position=para_num,
                    formatting=formatting
                )
                text_segments.append(segment)
                
                # Extract citations and academic terms
                citations.extend(self._extract_citations(paragraph.text))
                academic_terms.extend(self._extract_academic_terms(paragraph.text))
        
        except Exception as e:
            logger.error(f"DOCX structure extraction failed: {e}")
            # Fallback to basic extraction
            text_content = await self._extract_docx_text_basic(file_path)
            return text_content, DocumentStructure([], [], [], [], {})
        
        # Combine text segments
        full_text = '\n\n'.join([seg.text for seg in text_segments])
        
        # Create structure
        headings = [
            {"text": seg.text, "level": seg.level, "position": seg.position, "style": seg.style}
            for seg in text_segments if seg.segment_type == "heading"
        ]
        paragraphs = [
            {"text": seg.text, "position": seg.position, "formatting": seg.formatting}
            for seg in text_segments if seg.segment_type == "paragraph"
        ]
        
        structure = DocumentStructure(
            headings=headings,
            paragraphs=paragraphs,
            citations=list(set(citations)),
            academic_terms=list(set(academic_terms)),
            metadata={"total_paragraphs": len(text_segments)}
        )
        
        return full_text, structure
    
    async def _extract_txt_with_structure(self, file_path: Path) -> Tuple[str, DocumentStructure]:
        """Extract TXT file with basic structure detection."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as file:
                content = file.read()
        
        lines = content.split('\n')
        text_segments = []
        citations = []
        academic_terms = []
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            segment = TextSegment(
                text=line,
                segment_type=self._classify_text_segment(line),
                level=0,
                style=None,
                position=line_num,
                formatting={}
            )
            text_segments.append(segment)
            
            # Extract citations and academic terms
            citations.extend(self._extract_citations(line))
            academic_terms.extend(self._extract_academic_terms(line))
        
        # Create structure
        headings = [
            {"text": seg.text, "level": seg.level, "position": seg.position}
            for seg in text_segments if seg.segment_type == "heading"
        ]
        paragraphs = [
            {"text": seg.text, "position": seg.position}
            for seg in text_segments if seg.segment_type == "paragraph"
        ]
        
        structure = DocumentStructure(
            headings=headings,
            paragraphs=paragraphs,
            citations=list(set(citations)),
            academic_terms=list(set(academic_terms)),
            metadata={"total_lines": len(lines)}
        )
        
        return content, structure
    
    def _classify_text_segment(self, text: str) -> str:
        """Classify text segment type."""
        text = text.strip()
        
        # Check if it's a heading (common patterns)
        if (len(text) < 100 and 
            (text.isupper() or 
             re.match(r'^[IVX]+\.', text) or  # Roman numerals
             re.match(r'^\d+\.', text) or     # Numbers
             re.match(r'^[A-Z][a-z]*\s[A-Z]', text))):  # Title case
            return "heading"
        
        # Check if it's a citation or reference
        if any(re.search(pattern, text) for pattern in self.citation_patterns):
            return "citation"
        
        # Check if it's a list item
        if re.match(r'^[-•*]\s', text) or re.match(r'^\d+\.\s', text):
            return "list_item"
        
        # Default to paragraph
        return "paragraph"
    
    def _extract_citations(self, text: str) -> List[str]:
        """Extract citations from text."""
        citations = []
        for pattern in self.citation_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            citations.extend(matches)
        return citations
    
    def _extract_academic_terms(self, text: str) -> List[str]:
        """Extract academic terms from text."""
        academic_terms = []
        for pattern in self.academic_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            academic_terms.extend(matches)
        return academic_terms
    
    async def _preprocess_text(self, text: str, extract_academic_terms: bool = True) -> str:
        """Comprehensive text preprocessing."""
        # Clean text
        cleaned_text = self._clean_text(text)
        
        # Normalize whitespace
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        # Preserve academic terms if requested
        if extract_academic_terms:
            # This is a placeholder - in practice, you might want to
            # mark academic terms for protection during paraphrasing
            pass
        
        return cleaned_text.strip()
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common encoding issues
        text = text.replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"')
        text = text.replace('\u2013', '-').replace('\u2014', '-')
        
        # Remove unnecessary line breaks within sentences
        text = re.sub(r'(?<=[a-z])\n(?=[a-z])', ' ', text)
        
        return text
    
    def _calculate_quality_metrics(self, text: str) -> Dict[str, Any]:
        """Calculate text quality metrics."""
        try:
            metrics = {
                "readability_score": textstat.flesch_reading_ease(text),
                "grade_level": textstat.flesch_kincaid_grade(text),
                "sentence_count": textstat.sentence_count(text),
                "word_count": len(text.split()),
                "character_count": len(text),
                "avg_sentence_length": textstat.avg_sentence_length(text),
                "difficult_words": textstat.difficult_words(text)
            }
        except Exception as e:
            logger.warning(f"Failed to calculate quality metrics: {e}")
            metrics = {
                "word_count": len(text.split()),
                "character_count": len(text),
                "sentence_count": len(sent_tokenize(text)) if sent_tokenize else 0
            }
        
        return metrics
    
    async def _extract_text_basic(self, file_path: Path, file_ext: str) -> str:
        """Basic text extraction without structure preservation."""
        if file_ext == ".pdf":
            return await self._extract_pdf_text_basic(file_path)
        elif file_ext == ".docx":
            return await self._extract_docx_text_basic(file_path)
        elif file_ext == ".txt":
            return await self._extract_txt_text_basic(file_path)
        else:
            raise ValueError(f"Unsupported file extension: {file_ext}")
    
    async def _extract_pdf_text_basic(self, file_path: Path) -> str:
        """Basic PDF text extraction."""
        text_content = []
        
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
                    continue
        
        if not text_content:
            raise IOError("No text content found in PDF")
        
        return "\n\n".join(text_content)
    
    async def _extract_docx_text_basic(self, file_path: Path) -> str:
        """Basic DOCX text extraction."""
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            if not text_content:
                raise IOError("No text content found in DOCX")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise IOError(f"Failed to process DOCX file: {str(e)}")
    
    async def _extract_txt_text_basic(self, file_path: Path) -> str:
        """Basic TXT text extraction."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
            
            if not content.strip():
                raise IOError("Text file is empty")
            
            return content
            
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="latin-1") as file:
                    content = file.read()
                return content
            except Exception as e:
                raise IOError(f"Failed to decode text file: {str(e)}")
        except Exception as e:
            raise IOError(f"Failed to read text file: {str(e)}")
    
    async def reconstruct_document(
        self, 
        text: str, 
        original_structure: Optional[DocumentStructure] = None,
        output_path: Optional[Path] = None
    ) -> Path:
        """
        Reconstruct document from processed text maintaining original formatting.
        
        Args:
            text: Processed text content
            original_structure: Original document structure
            output_path: Output file path
            
        Returns:
            Path: Path to reconstructed document
        """
        if not output_path:
            output_path = self.upload_dir / f"reconstructed_{uuid.uuid4()}.docx"
        
        # Create new document
        doc = DocxDocument()
        
        if original_structure and original_structure.headings:
            # Reconstruct with original structure
            await self._reconstruct_with_structure(doc, text, original_structure)
        else:
            # Simple reconstruction
            await self._reconstruct_simple(doc, text)
        
        # Save document
        doc.save(output_path)
        
        logger.info(f"Document reconstructed: {output_path}")
        return output_path
    
    async def _reconstruct_with_structure(
        self, 
        doc: DocxDocument, 
        text: str, 
        structure: DocumentStructure
    ):
        """Reconstruct document with preserved structure."""
        # Split text into sentences/paragraphs
        sentences = sent_tokenize(text)
        sentence_idx = 0
        
        # Add headings and paragraphs in order
        for heading in structure.headings:
            # Add heading
            heading_para = doc.add_heading(heading["text"], level=heading.get("level", 1))
            
            # Add following paragraphs until next heading
            while sentence_idx < len(sentences):
                para = doc.add_paragraph(sentences[sentence_idx])
                sentence_idx += 1
                
                # Simple heuristic: stop if we've added enough content
                if sentence_idx % 3 == 0:  # Every 3 sentences, check for next heading
                    break
    
    async def _reconstruct_simple(self, doc: DocxDocument, text: str):
        """Simple document reconstruction."""
        paragraphs = text.split('\n\n')
        
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            
            # Check if it looks like a heading
            if len(para_text) < 100 and para_text[0].isupper():
                doc.add_heading(para_text, level=1)
            else:
                doc.add_paragraph(para_text)


# Global instance
enhanced_document_processor = EnhancedDocumentProcessor()
