"""
Document Processing Service
Handles document upload, text extraction, and content processing.
"""
import os
import uuid
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, BinaryIO

import docx
import PyPDF2
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document import Document, DocumentStatus
from app.core.database import get_db

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing uploaded documents."""
    
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
    
    async def upload_document(
        self, 
        file: BinaryIO, 
        filename: str,
        db: Session,
        chapter: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """
        Upload and process a document file.
        
        Args:
            file: File-like object containing the document
            filename: Original filename
            db: Database session
            chapter: Chapter designation (e.g., "BAB 1")
            metadata: Additional metadata
            
        Returns:
            Document: Created document record
            
        Raises:
            ValueError: If file type is not supported
            IOError: If file processing fails
        """
        # Validate file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in settings.allowed_extensions:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Generate unique file path
        file_id = str(uuid.uuid4())
        safe_filename = f"{file_id}{file_ext}"
        file_path = self.upload_dir / safe_filename
        
        try:
            # Save file to disk
            with open(file_path, "wb") as f:
                # Handle both async and sync file objects
                if hasattr(file, 'read') and callable(getattr(file, 'read', None)):
                    # Check if it's an async file
                    if hasattr(file, '__aiter__') or hasattr(file.read, '__await__'):
                        content = await file.read()
                    else:
                        content = file.read()
                else:
                    # If it's already bytes or other content
                    content = file
                f.write(content)
            
            # Extract text content
            text_content = await self._extract_text(file_path, file_ext)
            
            # Create document record
            document = Document(
                filename=filename,
                chapter=chapter,
                original_content=text_content,
                file_path=str(file_path),
                status=DocumentStatus.PENDING,
                document_metadata=metadata or {}
            )
            
            db.add(document)
            db.commit()
            db.refresh(document)
            
            logger.info(f"Document uploaded successfully: {document.id}")
            return document
            
        except Exception as e:
            # Clean up file if processing failed
            if file_path.exists():
                file_path.unlink()
            raise IOError(f"Failed to process document: {str(e)}")
    
    async def _extract_text(self, file_path: Path, file_ext: str) -> str:
        """
        Extract text content from uploaded file.
        
        Args:
            file_path: Path to the uploaded file
            file_ext: File extension
            
        Returns:
            str: Extracted text content
            
        Raises:
            ValueError: If file type is not supported
            IOError: If text extraction fails
        """
        try:
            if file_ext == ".pdf":
                return await self._extract_pdf_text(file_path)
            elif file_ext == ".docx":
                return await self._extract_docx_text(file_path)
            elif file_ext == ".txt":
                return await self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file extension: {file_ext}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            raise IOError(f"Failed to extract text: {str(e)}")
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text_content = []
        
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_content.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
                    continue
        
        if not text_content:
            raise IOError("No text content found in PDF")
        
        return "\n\n".join(text_content)
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            if not text_content:
                raise IOError("No text content found in DOCX")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            raise IOError(f"Failed to process DOCX file: {str(e)}")
    
    async def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
            
            if not content.strip():
                raise IOError("Text file is empty")
            
            return content
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, "r", encoding="latin-1") as file:
                    content = file.read()
                return content
            except Exception as e:
                raise IOError(f"Failed to decode text file: {str(e)}")
        except Exception as e:
            raise IOError(f"Failed to read text file: {str(e)}")
    
    def get_document(self, document_id: uuid.UUID, db: Session) -> Optional[Document]:
        """
        Retrieve document by ID.
        
        Args:
            document_id: Document UUID
            db: Database session
            
        Returns:
            Document or None if not found
        """
        return db.query(Document).filter(Document.id == document_id).first()
    
    def update_document_status(
        self, 
        document_id: uuid.UUID, 
        status: DocumentStatus,
        db: Session,
        paraphrased_content: Optional[str] = None
    ) -> Optional[Document]:
        """
        Update document processing status.
        
        Args:
            document_id: Document UUID
            status: New status
            db: Database session
            paraphrased_content: Paraphrased text if completed
            
        Returns:
            Updated Document or None if not found
        """
        document = self.get_document(document_id, db)
        if not document:
            return None
        
        document.status = status
        
        if status == DocumentStatus.COMPLETED:
            document.processed_date = datetime.utcnow()
            if paraphrased_content:
                document.paraphrased_content = paraphrased_content
        
        db.commit()
        db.refresh(document)
        
        logger.info(f"Document {document_id} status updated to {status}")
        return document
    
    def delete_document(self, document_id: uuid.UUID, db: Session) -> bool:
        """
        Delete document and associated file.
        
        Args:
            document_id: Document UUID
            db: Database session
            
        Returns:
            bool: True if deleted successfully
        """
        document = self.get_document(document_id, db)
        if not document:
            return False
        
        # Delete file from disk
        try:
            file_path = Path(document.file_path)
            if file_path.exists():
                file_path.unlink()
        except Exception as e:
            logger.warning(f"Failed to delete file {document.file_path}: {str(e)}")
        
        # Delete from database
        db.delete(document)
        db.commit()
        
        logger.info(f"Document {document_id} deleted successfully")
        return True


# Global instance
document_processor = DocumentProcessor()
